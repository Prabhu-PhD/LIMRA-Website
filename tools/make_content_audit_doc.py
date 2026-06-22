# Generates the LIMRA content audit as a Word document.
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x25, 0x24, 0x5D)
RED = RGBColor(0xC0, 0x24, 0x1D)
GREY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# Base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), hex_color)
    tcPr.append(sh)

def set_cell_text(cell, text, bold=False, color=None, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    runs = text.split("\n")
    for i, line in enumerate(runs):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        r.bold = bold
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
    return cell

def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=WHITE, size=9.5)
        shade(hdr[i], "25245D")
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=9.5)
        if ri % 2 == 1:
            for c in cells:
                shade(c, "F3F3F7")
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = NAVY

def h2(text, space_before=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RED

def body(text, italic=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text).font.size = Pt(10.5)

# ---------------- Title ----------------
h1("LIMRA Website — Content Audit")
sub = doc.add_paragraph()
r = sub.add_run("Placeholder vs. source-file content, for the client content update")
r.italic = True
r.font.color.rgb = GREY
r.font.size = Pt(10.5)
meta = doc.add_paragraph()
r = meta.add_run("Prepared 19 June 2026")
r.font.size = Pt(9)
r.font.color.rgb = GREY

# ---------------- Context ----------------
h2("How to read this", space_before=8)
body('The "source files" folder is data/ (six JSON files). Every .html page is hand-written static '
     "HTML (only the blog is templated), so the data files are not auto-injected — page content is "
     "hardcoded, and each block below is traced back to its origin. Key nuance: some source files are "
     'themselves placeholder, so "from a source file" does not always mean "real, final content."')

p = doc.add_paragraph()
for label, desc in [
    ("✅ Real", " — client/factual content; verify only."),
    ("❌ Placeholder", " — invented dummy content; must be replaced."),
    ("✍ Drafted copy", " — written for the site (not lorem, not from a source file); client should review."),
    ("⚠ Template/Draft", " — needs review/approval (e.g. legal, blog)."),
]:
    rr = p.add_run(label); rr.bold = True; rr.font.size = Pt(10)
    rr2 = p.add_run(desc + "   "); rr2.font.size = Pt(10)

# ---------------- Table A ----------------
h2("Table A — The source files (data/): real vs. placeholder")
add_table(
    ["Source file", "Feeds", "Status", "Notes"],
    [
        ["contact.json", "Footer (all pages), Contact page", "✅ Real",
         "Address, 5 phones, WhatsApp, email, website, director “Mohammed Ghani”. mapEmbed is empty, but the Contact page hardcodes a real Google Maps embed."],
        ["stats.json", "Hero stats, FMGE page, About, etc.", "✅ Real (client figures)",
         "2000+ doctors · 24 yrs · 92% FMGE · 50+ mocks · 800 hrs · 10 countries · 21 scored >200. See discrepancy note."],
        ["colleges.json", "Colleges list, college detail pages, Compare, home cards", "✅ Real",
         "6 real universities — names, cities, countries, programs, hospitals, badges. Descriptions are marketing copy. (JSON image paths are stale .jpg; pages use real .webp/.png.)"],
        ["testimonials.json", "Home carousel + Testimonials page", "❌ Placeholder",
         "6 invented “Dr.” names + generic quotes + scores. Photos empty."],
        ["gallery.json", "Nothing — unused", "❌ Placeholder & stale",
         "All 6 entries point to one image. The live gallery uses real photos instead, so this file is dead."],
        ["blog.json + blog/posts/*.md", "News / Blog", "⚠ Draft",
         "6 SEO articles authored by “LIMRA Team” — written for the site, not client-supplied."],
    ],
    widths=[1.5, 1.8, 1.2, 3.0],
)

# ---------------- Table B ----------------
h2("Table B — Site content by area: origin & action")
add_table(
    ["Page / Section", "Content", "Origin", "Client action"],
    [
        ["All footers + Contact", "Address, phones, email, WhatsApp, map", "contact.json ✅", "Verify only"],
        ["Everywhere", "Headline numbers (2000+, 24 yrs, 92%, 800 hrs…)", "stats.json ✅", "Verify (see ⚠)"],
        ["Colleges / detail / Compare / home cards", "6 universities, logos, programs, hospitals", "colleges.json ✅ + real logo assets", "Verify facts"],
        ["Gallery page + home photo strip", "Campus / graduation / facility photos", "Real curated /assets/*.webp", "Confirm these are actual LIMRA photos"],
        ["Contact", "Google Maps location", "Real embed", "Confirm the pin is correct"],
        ["Home hero — student orbit", "9 face photos", "❌ Placeholder (default-avatar.svg)", "Send 9 real student photos"],
        ["Testimonials (home + page)", "6 students, quotes, scores", "❌ Placeholder (testimonials.json)", "Replace with real testimonials"],
        ["About — Leadership", "4 cards", "Mixed: card 1 (Mohammed Ghani) ✅ real; cards 2–4 marked PLACEHOLDER “Team member”", "Add 3 real leaders or remove cards"],
        ["News / Blog", "6 articles", "⚠ Draft (AI-written)", "Review / approve or rewrite"],
        ["About — story + timeline", "2001→2026 milestones", "✍ Drafted copy", "Confirm dates/events are accurate"],
        ["About — Impact / advocacy band", "before→after figures", "✍ Drafted / illustrative", "Confirm the numbers"],
        ["Every page", "Hero headlines & taglines", "✍ Drafted copy", "Review wording"],
        ["Coaching / Medical Tourism / FAQ", "Body copy, Q&As, value props", "✍ Drafted copy", "Review for accuracy"],
        ["Privacy / Terms", "Legal text", "⚠ Template", "Lawyer review before relying on them"],
    ],
    widths=[1.9, 1.9, 1.9, 1.8],
)

# ---------------- Shortlist ----------------
h2("Must-replace before it counts as “real” content")
for s in [
    "Student orbit photos (home hero) — 9 real photos",
    "Testimonials — real students, quotes, results (+ optional photos)",
    "About leadership — 3 placeholder “Team member” cards",
    "Blog / News — approve or rewrite the 6 drafted articles",
    "Confirm gallery photos are genuinely LIMRA’s",
]:
    bullet(s)

# ---------------- Discrepancy ----------------
h2("Discrepancy to settle with the client")
body('The Lyceum brochure says “10,000+ doctors” but the whole site says “2000+”. '
     "Pick the correct figure — it lives in stats.json and is hardcoded across many pages.")

# ---------------- Footer note ----------------
note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(16)
r = note.add_run("Tip: I can wire these pages to read from data/ so the client edits the JSON in one "
                 "place instead of hunting through the HTML — ask if you’d like that.")
r.italic = True
r.font.size = Pt(9.5)
r.font.color.rgb = GREY

out = "LIMRA-Content-Audit.docx"
doc.save(out)
print("saved", out)
